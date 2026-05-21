#!/usr/bin/env python3
"""
generate_docx.py — Generate a fresh, ATS-safe DOCX from structured resume content.

Supports two templates:
  classic    — Times New Roman, centered header, horizontal-rule section headers
  executive  — Calibri/sans-serif, left-aligned header with right-aligned contact

Input (stdin, JSON):
{
  "content": {
    "header": { "name":"", "email":"", "phone":"", "location":"", "linkedin":"", "github":"" },
    "summary": "",
    "skills": { "Cloud Platforms": ["AWS","Azure"], ... },
    "experiences": [
      { "company":"", "role":"", "dates":"", "location":"", "client":"", "bullets":[] }
    ],
    "education": [{ "school":"", "degree":"", "start":"", "end":"", "gpa":"", "location":"" }],
    "certifications": [{ "name":"", "date":"" }]
  },
  "selectedPointsByRole": [
    { "roleName":"", "company":"", "dates":"", "bullets":[] }
  ],
  "templateName": "classic"
}

Output (stdout): base64-encoded DOCX
"""

import sys, json, base64, re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ────────────────────────────────────────────────

def set_font(run, name, size_pt, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_hr(para, color=(17, 17, 17), width_pt=1):
    """Add a bottom border (horizontal rule) to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'),  str(int(width_pt * 8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(*color))
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_paragraph_spacing(para, before_pt=0, after_pt=2, line_rule=None, line_val=None):
    pPr  = para._p.get_or_add_pPr()
    pSpc = OxmlElement('w:spacing')
    pSpc.set(qn('w:before'), str(int(before_pt * 20)))
    pSpc.set(qn('w:after'),  str(int(after_pt  * 20)))
    if line_rule and line_val:
        pSpc.set(qn('w:lineRule'), line_rule)
        pSpc.set(qn('w:line'),     str(line_val))
    pPr.append(pSpc)

def set_margins(doc, top=0.75, bottom=0.75, left=0.75, right=0.75):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)

def norm(s):
    return (s or '').lower().replace('\t', ' ').strip()


# ── Merge selected bullets into experiences ────────────────

def merge_bullets(experiences, selected_by_role):
    """
    Merge optimizer-selected bullets into the matching experience entry.
    Matches by company name (case-insensitive, first 12 chars).
    Appends selected bullets after existing ones.
    """
    if not selected_by_role:
        return experiences

    result = [dict(e) for e in experiences]

    for sel in selected_by_role:
        if not sel.get('bullets'):
            continue
        sel_key = re.sub(r'[^a-z0-9]', '', norm(sel.get('company', sel.get('roleName', ''))))[:12]
        matched = False
        for exp in result:
            exp_key = re.sub(r'[^a-z0-9]', '', norm(exp.get('company', '')))[:12]
            if sel_key and exp_key and (sel_key in exp_key or exp_key in sel_key):
                existing = list(exp.get('bullets') or [])
                for b in (sel.get('bullets') or []):
                    if b and b not in existing:
                        existing.append(b)
                exp['bullets'] = existing
                matched = True
                break
        # If no match by company, try by role index
        if not matched and result:
            idx = selected_by_role.index(sel) if sel in selected_by_role else 0
            if idx < len(result):
                existing = list(result[idx].get('bullets') or [])
                for b in (sel.get('bullets') or []):
                    if b and b not in existing:
                        existing.append(b)
                result[idx]['bullets'] = existing

    return result


# ── Classic Professional Template ─────────────────────────
# Times New Roman, centered header, uppercase section headers with bottom border

def build_classic(doc, content):
    FONT     = 'Times New Roman'
    COLOR_BK = (17, 17, 17)      # near-black

    header = content.get('header') or {}
    summ   = (content.get('summary') or '').strip()
    skills = content.get('skills') or {}
    exps   = content.get('experiences') or []
    edu    = content.get('education') or []
    certs  = content.get('certifications') or []

    def section_header(text):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=6, after_pt=3)
        run = p.add_run(text.upper())
        set_font(run, FONT, 11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_hr(p, color=COLOR_BK, width_pt=1.5)
        return p

    def body_para(text='', indent_pt=0):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=2)
        if text:
            run = p.add_run(text)
            set_font(run, FONT, 10)
        if indent_pt:
            p.paragraph_format.left_indent = Pt(indent_pt)
        return p

    # ── Header ──────────────────────────────────────────────
    name_para = doc.add_paragraph()
    set_paragraph_spacing(name_para, before_pt=0, after_pt=2)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(header.get('name') or 'Your Name')
    set_font(name_run, FONT, 18, bold=True)

    contact_parts = [p for p in [
        header.get('phone'), header.get('email'), header.get('location')
    ] if p]
    if contact_parts:
        cp = doc.add_paragraph()
        set_paragraph_spacing(cp, before_pt=0, after_pt=1)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run('  |  '.join(contact_parts))
        set_font(cr, FONT, 10)

    link_parts = [p for p in [header.get('linkedin'), header.get('github')] if p]
    if link_parts:
        lp = doc.add_paragraph()
        set_paragraph_spacing(lp, before_pt=0, after_pt=4)
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run('  |  '.join(link_parts))
        set_font(lr, FONT, 10, color=(85, 85, 85))

    # ── Summary ─────────────────────────────────────────────
    if summ:
        section_header('Professional Summary')
        sp = body_para()
        set_paragraph_spacing(sp, before_pt=0, after_pt=4)
        sr = sp.add_run(summ)
        set_font(sr, FONT, 10)
        sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Technical Skills ────────────────────────────────────
    flat_skills = {k: v for k, v in skills.items() if v}
    if flat_skills:
        section_header('Technical Skills')
        for cat, items in flat_skills.items():
            vals = items if isinstance(items, list) else [items]
            p = body_para()
            set_paragraph_spacing(p, before_pt=0, after_pt=2)
            br = p.add_run(cat + ': ')
            set_font(br, FONT, 10, bold=True)
            vr = p.add_run(', '.join(str(v) for v in vals))
            set_font(vr, FONT, 10)

    # ── Professional Experience ──────────────────────────────
    if exps:
        section_header('Professional Experience')
        for exp in exps:
            comp  = (exp.get('company') or '').strip()
            role  = (exp.get('role')    or '').strip()
            dates = (exp.get('dates')   or '').strip()
            loc   = (exp.get('location') or '').strip()
            cli   = (exp.get('client')   or '').strip()
            bulls = exp.get('bullets') or []

            # Company | Role line with date right-aligned via tab
            ep = doc.add_paragraph()
            set_paragraph_spacing(ep, before_pt=4, after_pt=1)
            # Set tab stop at right margin
            from docx.oxml import OxmlElement as OE
            pPr = ep._p.get_or_add_pPr()
            tabs = OE('w:tabs')
            tab  = OE('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), '9072')   # ~6.3in at 1440 twips/inch
            tabs.append(tab)
            pPr.append(tabs)

            cr = ep.add_run(comp)
            set_font(cr, FONT, 11, bold=True)
            if role:
                pr = ep.add_run(' | ')
                set_font(pr, FONT, 10)
                rr = ep.add_run(role)
                set_font(rr, FONT, 10, italic=True)
            if dates:
                dt = ep.add_run('\t' + dates)
                set_font(dt, FONT, 10)

            # Location / Client sub-line
            if loc or cli:
                sub_parts = [p for p in [cli, loc] if p]
                sp2 = doc.add_paragraph()
                set_paragraph_spacing(sp2, before_pt=0, after_pt=1)
                sr2 = sp2.add_run('  '.join(sub_parts))
                set_font(sr2, FONT, 9.5, italic=True, color=(85, 85, 85))

            # Bullets
            for b in bulls:
                if not b or not b.strip():
                    continue
                bp = doc.add_paragraph()
                set_paragraph_spacing(bp, before_pt=0, after_pt=1)
                bp.paragraph_format.left_indent  = Pt(12)
                bp.paragraph_format.first_line_indent = Pt(-12)
                br = bp.add_run('\u2022 ' + b.strip())
                set_font(br, FONT, 10)

    # ── Education ───────────────────────────────────────────
    if edu:
        section_header('Education')
        for e in edu:
            school = (e.get('school') or '').strip()
            degree = (e.get('degree') or '').strip()
            start  = (e.get('start')  or '').strip()
            end    = (e.get('end')    or '').strip()
            gpa    = (e.get('gpa')    or '').strip()
            eloc   = (e.get('location') or '').strip()

            ep = doc.add_paragraph()
            set_paragraph_spacing(ep, before_pt=3, after_pt=1)
            sr = ep.add_run(school)
            set_font(sr, FONT, 10, bold=True)
            if degree:
                dr = ep.add_run(' | ' + degree)
                set_font(dr, FONT, 10)
            date_str = ' - '.join(filter(None, [start, end]))
            if date_str:
                dtr = ep.add_run('\t' + date_str)
                set_font(dtr, FONT, 10)
            if gpa or eloc:
                sp2 = doc.add_paragraph()
                set_paragraph_spacing(sp2, before_pt=0, after_pt=1)
                sr2 = sp2.add_run('  '.join(filter(None, [eloc, ('GPA: ' + gpa) if gpa else ''])))
                set_font(sr2, FONT, 9.5, italic=True, color=(85, 85, 85))

    # ── Certifications ──────────────────────────────────────
    if certs:
        section_header('Certifications')
        # Two per row using tab
        for i in range(0, len(certs), 2):
            cp = doc.add_paragraph()
            set_paragraph_spacing(cp, before_pt=0, after_pt=2)
            name1 = (certs[i].get('name') or '')
            date1 = (certs[i].get('date') or '')
            text1 = name1 + (' (' + date1 + ')' if date1 else '')
            cr1 = cp.add_run('\u2022 ' + text1)
            set_font(cr1, FONT, 10)
            if i + 1 < len(certs):
                name2 = (certs[i+1].get('name') or '')
                date2 = (certs[i+1].get('date') or '')
                text2 = name2 + (' (' + date2 + ')' if date2 else '')
                cr2 = cp.add_run('\t\u2022 ' + text2)
                set_font(cr2, FONT, 10)


# ── Executive Compact Template ────────────────────────────
# Calibri/Segoe, left-aligned name with right-aligned contact, dense layout

def build_executive(doc, content):
    FONT     = 'Calibri'
    COLOR_DK = (13, 27, 42)    # #0D1B2A

    header = content.get('header') or {}
    summ   = (content.get('summary') or '').strip()
    skills = content.get('skills') or {}
    exps   = content.get('experiences') or []
    edu    = content.get('education') or []
    certs  = content.get('certifications') or []

    def section_header(text):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=6, after_pt=2)
        run = p.add_run(text.upper())
        set_font(run, FONT, 9.5, bold=True, color=COLOR_DK)
        add_hr(p, color=COLOR_DK, width_pt=2)
        return p

    # ── Header ──────────────────────────────────────────────
    hp = doc.add_paragraph()
    set_paragraph_spacing(hp, before_pt=0, after_pt=2)
    # Name left, contact right via tab
    nr = hp.add_run(header.get('name') or 'Your Name')
    set_font(nr, FONT, 18, bold=True, color=COLOR_DK)
    contact_parts = [p for p in [
        header.get('email'), header.get('phone'), header.get('location')
    ] if p]
    if contact_parts:
        ct = hp.add_run('\t' + '  \u00b7  '.join(contact_parts))
        set_font(ct, FONT, 9, color=(85, 85, 85))
    add_hr(hp, color=COLOR_DK, width_pt=2)

    link_parts = [p for p in [header.get('linkedin'), header.get('github')] if p]
    if link_parts:
        lp = doc.add_paragraph()
        set_paragraph_spacing(lp, before_pt=0, after_pt=4)
        lr = lp.add_run('  \u00b7  '.join(link_parts))
        set_font(lr, FONT, 9, color=(85, 85, 85))

    # ── Summary ─────────────────────────────────────────────
    if summ:
        section_header('Summary')
        sp = doc.add_paragraph()
        set_paragraph_spacing(sp, before_pt=0, after_pt=4)
        sr = sp.add_run(summ)
        set_font(sr, FONT, 10)

    # ── Technical Skills ────────────────────────────────────
    flat_skills = {k: v for k, v in skills.items() if v}
    if flat_skills:
        section_header('Technical Skills')
        for cat, items in flat_skills.items():
            vals = items if isinstance(items, list) else [items]
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before_pt=0, after_pt=1)
            br = p.add_run(cat + ': ')
            set_font(br, FONT, 9.5, bold=True, color=COLOR_DK)
            vr = p.add_run('  \u00b7  '.join(str(v) for v in vals))
            set_font(vr, FONT, 9.5)

    # ── Experience ───────────────────────────────────────────
    if exps:
        section_header('Experience')
        for exp in exps:
            comp  = (exp.get('company') or '').strip()
            role  = (exp.get('role')    or '').strip()
            dates = (exp.get('dates')   or '').strip()
            cli   = (exp.get('client')  or '').strip()
            bulls = exp.get('bullets') or []

            ep = doc.add_paragraph()
            set_paragraph_spacing(ep, before_pt=4, after_pt=0)
            cr = ep.add_run(comp)
            set_font(cr, FONT, 10, bold=True, color=COLOR_DK)
            if role:
                rr = ep.add_run(' - ' + role)
                set_font(rr, FONT, 10, italic=True)
            if dates:
                dt = ep.add_run('\t' + dates)
                set_font(dt, FONT, 9.5, color=(100, 100, 100))
            if cli:
                cp2 = doc.add_paragraph()
                set_paragraph_spacing(cp2, before_pt=0, after_pt=0)
                cr2 = cp2.add_run('Client: ' + cli)
                set_font(cr2, FONT, 9.5, italic=True, color=(85, 85, 85))

            for b in bulls:
                if not b or not b.strip():
                    continue
                bp = doc.add_paragraph()
                set_paragraph_spacing(bp, before_pt=0, after_pt=1)
                bp.paragraph_format.left_indent       = Pt(10)
                bp.paragraph_format.first_line_indent = Pt(-10)
                br = bp.add_run('\u2022 ' + b.strip())
                set_font(br, FONT, 9.5)

    # ── Education ───────────────────────────────────────────
    if edu:
        section_header('Education')
        for e in edu:
            school = (e.get('school') or '').strip()
            degree = (e.get('degree') or '').strip()
            start  = (e.get('start')  or '').strip()
            end    = (e.get('end')    or '').strip()
            ep = doc.add_paragraph()
            set_paragraph_spacing(ep, before_pt=2, after_pt=1)
            sr = ep.add_run(school)
            set_font(sr, FONT, 10, bold=True)
            if degree:
                dr = ep.add_run('  \u00b7  ' + degree)
                set_font(dr, FONT, 9.5)
            date_str = ' - '.join(filter(None, [start, end]))
            if date_str:
                dtr = ep.add_run('\t' + date_str)
                set_font(dtr, FONT, 9.5, color=(100, 100, 100))

    # ── Certifications ──────────────────────────────────────
    if certs:
        section_header('Certifications')
        cert_text = '  \u00b7  '.join(
            (c.get('name') or '') + (' (' + c.get('date') + ')' if c.get('date') else '')
            for c in certs
        )
        cp = doc.add_paragraph()
        set_paragraph_spacing(cp, before_pt=0, after_pt=2)
        cr = cp.add_run(cert_text)
        set_font(cr, FONT, 9.5)


# ── Main ──────────────────────────────────────────────────

def main():
    data     = json.loads(sys.stdin.read())
    content  = data.get('content') or {}
    selected = data.get('selectedPointsByRole') or []
    template = (data.get('templateName') or 'classic').lower()

    # Merge optimizer bullets into experiences
    content = dict(content)
    content['experiences'] = merge_bullets(
        content.get('experiences') or [],
        selected
    )

    doc = Document()
    set_margins(doc, top=0.75, bottom=0.75, left=0.75, right=0.75)

    # Remove default empty paragraph Word always adds
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    if template == 'executive':
        build_executive(doc, content)
    else:
        build_classic(doc, content)

    out = BytesIO()
    doc.save(out)
    sys.stdout.buffer.write(base64.b64encode(out.getvalue()))


if __name__ == '__main__':
    main()
