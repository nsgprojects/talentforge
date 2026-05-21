#!/usr/bin/env python3
"""
insert_bullets.py — Table-aware DOCX bullet inserter.

Stage 0 (NEW, PRIMARY for table resumes):
  Build an ordered list of "experience tables" — tables that contain a date cell.
  Select the Nth one by role_index. This is 100% positional and bypasses
  company-name matching entirely (which fails when names contain locations).

Stage 1: Match in Client: paragraph lines
Stage 2: Match company name in table cell text
Stage 3: Broad paragraph search
Stage 4: Positional date-heuristic fallback
"""

import sys, json, base64, random, re
from io import BytesIO
from copy import deepcopy
from collections import defaultdict
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

YEAR_RE  = re.compile(r'\b(19|20)\d{2}\b')
DATE_RE  = re.compile(r'\b(present|current|now)\b', re.I)


def norm(s):
    return (s or '').lower() \
        .replace('\t',' ').replace('\xa0',' ') \
        .replace('\u2013','-').replace('\u2014','-').strip()

def norm_search(s):
    return re.sub(r'[^a-z0-9]', '', norm(s))


def build_element_map(doc):
    """
    Ordered list of body elements with para indices and table metadata.
    For each TABLE, records:
      - first_para_after: para index of first real body para following the table
      - is_experience_table: True if any cell contains a date pattern
      - cell_texts: all cell text strings
    """
    raw = []
    para_count = table_count = 0

    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            text = ''.join(n.text or '' for n in child.iter() if n.tag == qn('w:t'))
            raw.append({'type':'para','idx':para_count,'text':text,'elem':child})
            para_count += 1
        elif tag == 'tbl':
            cell_texts = []
            for row in child.findall('.//' + qn('w:tr')):
                for cell in row.findall('.//' + qn('w:tc')):
                    # Get per-paragraph text within the cell
                    paras = []
                    for cp in cell.findall('.//' + qn('w:p')):
                        pt = ''.join(n.text or '' for n in cp.iter() if n.tag == qn('w:t')).strip()
                        if pt: paras.append(pt)
                    ct = ' | '.join(paras)
                    if ct: cell_texts.append(ct)
            full_text = ' '.join(cell_texts)
            is_exp = (YEAR_RE.search(full_text) or DATE_RE.search(full_text)) and len(full_text) < 200
            raw.append({
                'type':'table','idx':table_count,
                'text':full_text,'cell_texts':cell_texts,
                'is_experience_table': bool(is_exp),
                'elem':child
            })
            table_count += 1

    # Attach first_para_after to each table
    for i, el in enumerate(raw):
        if el['type'] == 'table':
            first_after = -1
            for j in range(i+1, len(raw)):
                if raw[j]['type'] == 'para' and raw[j]['text'].strip():
                    first_after = raw[j]['idx']
                    break
            el['first_para_after'] = first_after

    return raw


def get_experience_tables(element_map):
    """
    Returns element_map entries for experience tables in document order.
    An experience table is one that has a date cell (short cell with year/present).
    """
    return [el for el in element_map if el['type'] == 'table' and el.get('is_experience_table')]


def all_client_para_indices(paragraphs):
    result = []
    for i, p in enumerate(paragraphs):
        if re.match(r'^(client|employer|company)\s*[:\-]', norm(p.text)):
            result.append(i)
    return result


def find_experience_para_indices(paragraphs):
    skip = ('experience','education','skills','summary','certif',
            'project','objective','profile','technical','earlier','core')
    result = []
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if not t or len(t) > 150: continue
        if any(norm(t).startswith(s) for s in skip): continue
        if t[0] in '-•*·◦▪▸': continue
        if YEAR_RE.search(t):
            result.append(i)
    return result


def find_insertion_context(paragraphs, element_map, exp_tables,
                           company, role_index,
                           client_para_idx, experience_para_idx):
    """
    Stage 0 (PRIMARY): Use ordered experience tables — select Nth by role_index.
    Stage 1: Match in Client: paragraph lines.
    Stage 2: Match company name in table cell text.
    Stage 3: Broad paragraph search.
    Stage 4: Positional date-heuristic.
    """
    key = norm_search(company)
    bad_words = {'current','employer','previous','company','unnamed','recent',
                 'specified','ltd','limited','pvt','inc','corp','az','ca','tx',
                 'ny','ma','il','wa','fl','ga','co','oh','mi','pa','phoenix',
                 'boston','chicago','dallas','seattle','atlanta','denver','houston'}
    key_words = set(re.sub(r'[^a-z]',' ', norm(company)).split())
    is_generic = len(key) < 3 or not (key_words - bad_words)

    # ── Stage 0: positional experience table (PRIMARY) ───────
    if exp_tables and role_index < len(exp_tables):
        fpa = exp_tables[role_index].get('first_para_after', -1)
        if fpa >= 0:
            sys.stderr.write(
                f"INFO: Stage 0 positional table role_index={role_index} "
                f"→ first_para_after={fpa}: {paragraphs[fpa].text[:50]!r}\n"
            )
            return fpa, 'after_table'

    if not is_generic:
        # ── Stage 1: Client: paragraph ───────────────────────
        for pi in client_para_idx:
            if key[:10] in norm_search(paragraphs[pi].text):
                sys.stderr.write(f"INFO: Stage 1 Client: match for {company!r} → para {pi}\n")
                return pi, 'after_para'

        # ── Stage 2: table cell match ─────────────────────────
        for el in element_map:
            if el['type'] == 'table' and key[:10] in norm_search(el['text']):
                fpa = el.get('first_para_after', -1)
                if fpa >= 0:
                    sys.stderr.write(
                        f"INFO: Stage 2 table-match for {company!r} → first_para_after={fpa}\n"
                    )
                    return fpa, 'after_table'

        # ── Stage 3: broad paragraph search ──────────────────
        for i, p in enumerate(paragraphs):
            if key[:10] in norm_search(p.text) and len(p.text.strip()) < 200:
                sys.stderr.write(f"INFO: Stage 3 broad match for {company!r} at para {i}\n")
                return i, 'after_para'

    # ── Stage 4: positional date-heuristic ───────────────────
    if client_para_idx and role_index < len(client_para_idx):
        pi = client_para_idx[role_index]
        sys.stderr.write(f"INFO: Stage 4a Client: list role_index={role_index} → para {pi}\n")
        return pi, 'after_para'

    if experience_para_idx and role_index < len(experience_para_idx):
        pi = experience_para_idx[role_index]
        sys.stderr.write(f"INFO: Stage 4b date-heuristic role_index={role_index} → para {pi}\n")
        return pi, 'after_para'

    sys.stderr.write(f"WARNING: No anchor found for {company!r} role_index={role_index}\n")
    return -1, 'after_para'


def find_responsibilities_start(paragraphs, anchor_idx, section_type):
    if section_type == 'after_table':
        return anchor_idx
    for i in range(anchor_idx, min(anchor_idx+8, len(paragraphs))):
        t = norm(paragraphs[i].text)
        if ('roles' in t and 'responsib' in t) or 'responsibilities' in t:
            return i
    return anchor_idx + 2


def find_section_end(paragraphs, element_map, exp_tables, resp_idx, role_index):
    """
    For table-format resumes: section ends at the start of the NEXT experience table.
    For paragraph-format: use Environment: or next Client: line.
    """
    # If we have experience tables and the next role has a table, use that as the boundary
    if exp_tables and role_index + 1 < len(exp_tables):
        next_table_fpa = exp_tables[role_index + 1].get('first_para_after', -1)
        if next_table_fpa > resp_idx:
            sys.stderr.write(
                f"INFO: Section end for role_index={role_index} → next table fpa={next_table_fpa}\n"
            )
            return next_table_fpa  # stop just before next role starts

    # Fallback: scan for Environment: or next Client: line
    limit = min(resp_idx + 120, len(paragraphs))
    for i in range(resp_idx + 1, limit):
        t = norm(paragraphs[i].text)
        raw = paragraphs[i].text.strip()
        if t.startswith('environment:'):
            return i
        if re.match(r'^(client|employer|company)\s*[:\-]', t) and i > resp_idx + 2:
            return i
        if raw and raw == raw.upper() and 3 < len(raw) < 50 and i > resp_idx + 3:
            return i
    return min(resp_idx + 50, len(paragraphs))


def get_content_indices(paragraphs, resp_idx, end_idx):
    return [i for i in range(resp_idx + 1, end_idx) if paragraphs[i].text.strip()]


def clone_paragraph(source_p_elem, new_text):
    new_p = deepcopy(source_p_elem)
    for tag in (qn('w:r'), qn('w:hyperlink')):
        for el in new_p.findall(tag):
            new_p.remove(el)

    # Copy run properties from first run of source
    source_rPr = None
    for r in source_p_elem.findall(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            source_rPr = deepcopy(rPr)
            break

    if source_rPr is not None:
        # Strip bold — inserted bullets should never be bold regardless of template
        for bold_tag in (qn('w:b'), qn('w:bCs')):
            for b_el in source_rPr.findall(bold_tag):
                source_rPr.remove(b_el)

    new_r = OxmlElement('w:r')
    if source_rPr is not None:
        new_r.append(source_rPr)
    new_t = OxmlElement('w:t')
    new_t.text = new_text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p


def insert_after_para(body, after_para_idx, new_p_elems):
    ref_p, count = None, 0
    for child in body:
        if child.tag == qn('w:p'):
            if count == after_para_idx:
                ref_p = child
                break
            count += 1
    if ref_p is None:
        return
    for new_p in reversed(new_p_elems):
        ref_p.addnext(new_p)


def main():
    data     = json.loads(sys.stdin.read())
    docx_b64 = data['docx']
    roles    = data['roles']

    doc        = Document(BytesIO(base64.b64decode(docx_b64)))
    paragraphs = doc.paragraphs
    body       = doc.element.body

    element_map         = build_element_map(doc)
    exp_tables          = get_experience_tables(element_map)
    client_para_idx     = all_client_para_indices(paragraphs)
    experience_para_idx = find_experience_para_indices(paragraphs) if not client_para_idx else []

    sys.stderr.write(
        f"INFO: {len(client_para_idx)} Client: sections {client_para_idx[:8]}\n"
        f"INFO: {len(exp_tables)} experience tables in document order\n"
    )
    for i, et in enumerate(exp_tables):
        sys.stderr.write(f"  exp_table[{i}]: first_para_after={et.get('first_para_after')} cells={et['cell_texts'][:2]}\n")

    # Resolve all insertions BEFORE modifying (merge duplicate sections)
    resolved = {}

    for role_data in roles:
        company    = role_data.get('company', '')
        bullets    = role_data.get('bullets', [])
        role_index = role_data.get('role_index', 0)
        if not bullets:
            continue

        anchor_idx, section_type = find_insertion_context(
            paragraphs, element_map, exp_tables,
            company, role_index, client_para_idx, experience_para_idx
        )
        if anchor_idx == -1:
            continue

        resp_idx     = find_responsibilities_start(paragraphs, anchor_idx, section_type)
        end_idx      = find_section_end(paragraphs, element_map, exp_tables, resp_idx, role_index)
        content_idxs = get_content_indices(paragraphs, resp_idx, end_idx)

        if not content_idxs:
            content_idxs = [resp_idx]

        sys.stderr.write(
            f"INFO: Inserting {len(bullets)} bullets for role_index={role_index} ({company!r}), "
            f"type={section_type}, content [{content_idxs[0]}..{content_idxs[-1]}]\n"
        )

        # Merge if same section hit by multiple roles
        if resp_idx not in resolved:
            resolved[resp_idx] = {'content_idxs': content_idxs, 'bullets': []}
        resolved[resp_idx]['bullets'].extend(bullets)

    # Insert bottom-up
    for resp_idx in sorted(resolved.keys(), reverse=True):
        sec          = resolved[resp_idx]
        content_idxs = sec['content_idxs']
        bullets      = sec['bullets']

        # Find best template paragraph:
        # Prefer List Paragraph style, non-bold, not an Environment: line
        template_p = None
        for idx in reversed(content_idxs):
            p   = paragraphs[idx]
            txt = p.text.strip().lower()
            if txt.startswith('environment:'):
                continue
            # Check rPr bold directly
            is_bold = any(
                r.find(qn('w:rPr')) is not None and
                r.find(qn('w:rPr')).find(qn('w:b')) is not None
                for r in p._p.findall(qn('w:r'))
            )
            if not is_bold:
                template_p = p._p
                break
        if template_p is None:
            # Fallback: use first content paragraph
            template_p = paragraphs[content_idxs[0]]._p

        random.shuffle(bullets)
        positions  = [random.choice(content_idxs) for _ in bullets]

        groups = defaultdict(list)
        for bullet, pos in zip(bullets, positions):
            groups[pos].append(bullet)

        for idx in sorted(groups.keys(), reverse=True):
            new_elems = [clone_paragraph(template_p, b) for b in groups[idx]]
            insert_after_para(body, idx, new_elems)
            paragraphs = doc.paragraphs

    out = BytesIO()
    doc.save(out)
    print(base64.b64encode(out.getvalue()).decode(), end='')


if __name__ == '__main__':
    main()
